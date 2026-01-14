"""DOCX file parser with formatting preservation."""

from pathlib import Path

from docx import Document
from docx.shared import Pt

from .base import BaseParser, TextSegment


class DocxParser(BaseParser):
    """Parser for DOCX files that preserves formatting."""

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx"]

    def parse(self, file_path: Path) -> list[TextSegment]:
        doc = Document(file_path)
        segments = []

        for para in doc.paragraphs:
            if not para.text.strip():
                continue

            # Store paragraph-level formatting
            metadata = {
                "style_name": para.style.name if para.style else None,
                "alignment": para.alignment,
                "runs": [],
            }

            # Capture run-level formatting (bold, italic, etc.)
            for run in para.runs:
                run_meta = {
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "font_name": run.font.name,
                    "font_size": run.font.size.pt if run.font.size else None,
                }
                metadata["runs"].append(run_meta)

            segments.append(
                TextSegment(
                    text=para.text,
                    metadata=metadata,
                    segment_type="paragraph",
                )
            )

        return segments

    def write(
        self,
        segments: list[TextSegment],
        output_path: Path,
        template_path: Path | None = None,
    ) -> None:
        # Use template to preserve document styles if provided
        if template_path:
            doc = Document(template_path)
            # Clear existing paragraphs but keep document structure
            for para in doc.paragraphs:
                para.clear()
            existing_paras = list(doc.paragraphs)
        else:
            doc = Document()
            existing_paras = []

        for i, segment in enumerate(segments):
            if i < len(existing_paras):
                para = existing_paras[i]
            else:
                para = doc.add_paragraph()

            # Apply paragraph-level formatting
            if segment.metadata.get("style_name"):
                try:
                    para.style = segment.metadata["style_name"]
                except KeyError:
                    pass  # Style not in document

            if segment.metadata.get("alignment") is not None:
                para.alignment = segment.metadata["alignment"]

            # Reconstruct runs with formatting
            runs_meta = segment.metadata.get("runs", [])
            if runs_meta:
                self._apply_run_formatting(para, segment.text, runs_meta)
            else:
                para.add_run(segment.text)

        doc.save(output_path)

    def _apply_run_formatting(
        self, para, translated_text: str, original_runs: list[dict]
    ) -> None:
        """
        Apply original run formatting to translated text.

        Since translation can change text length, we distribute formatting
        proportionally across the translated text.
        """
        if not original_runs:
            para.add_run(translated_text)
            return

        # Calculate original text length and proportions
        original_length = sum(len(run["text"]) for run in original_runs)
        if original_length == 0:
            para.add_run(translated_text)
            return

        translated_length = len(translated_text)
        current_pos = 0

        for run_meta in original_runs:
            # Calculate proportional length for this run
            proportion = len(run_meta["text"]) / original_length
            run_length = int(translated_length * proportion)

            # Handle last run - take remainder
            if run_meta == original_runs[-1]:
                run_text = translated_text[current_pos:]
            else:
                run_text = translated_text[current_pos : current_pos + run_length]

            if run_text:
                run = para.add_run(run_text)

                # Apply formatting
                if run_meta.get("bold"):
                    run.bold = True
                if run_meta.get("italic"):
                    run.italic = True
                if run_meta.get("underline"):
                    run.underline = True
                if run_meta.get("font_name"):
                    run.font.name = run_meta["font_name"]
                if run_meta.get("font_size"):
                    run.font.size = Pt(run_meta["font_size"])

            current_pos += run_length
